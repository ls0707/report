# -*- coding:utf-8 -*-

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.document import Document as _Document
from docx.table import Table, _Row, _Cell
from docx.text.paragraph import Paragraph
from config import Config
from docx.opc.exceptions import PackageNotFoundError
from clint.textui import prompt, puts, colored, validators
import re
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement


class DocxProcessor(object):
    # ----------------------------- Base Task
    class Task(object):
        def __init__(self):
            self.results = list()

        def _append_result(self):
            pass

    # ----------------------------- Replace Task
    class ParagraphReplaceTask(Task):
        def __init__(self, config):
            self.cfg = config
            super(DocxProcessor.ParagraphReplaceTask, self).__init__()

        def __call__(self, paragraph):
            for old, new in self.cfg.replace_list:
                if old in paragraph.text:
                    # for r in paragraph.runs:
                    #     if old in r.text:
                    paragraph.text = paragraph.text.replace(old, new)
            # print(paragraph.text)

    # ----------------------------- Chapter 4Task
    class ParagraphChapter4Task(Task):
        def __init__(self, config):
            self.cfg = config

            super(DocxProcessor.ParagraphChapter4Task, self).__init__()
            self.is_chapter4 = False  # 第四章起始
            self.is_after_chapter4 = False  # 第四章结束
            self.load_table_flag = False
            self.cycle_flag = False
            self.step_count = -1  # 循环序号
            self.section_name = ''
            self.paragraph_ro1 = re.compile('(?<=表)4-\d{1,2}-[12]')
            self.paragraph_ro2 = re.compile('(结果汇总)|(结果分析)')

            self.table_col_names = list()
            self.table_satisfy_values = list()
            self.table_some_and_not_satisfy_values = list()
            self.unsatisfied_table_id = ''

        def __call__(self, paragraph):
            if not self.is_after_chapter4:  # 遇到"单元测评小结"（单元测评结束）之前前始终检查每个段落
                if not self.is_chapter4:
                    if paragraph.text == '单元测评':
                        self.is_chapter4 = True
                else:  # 此处为开始处理文档内容的部分
                    if paragraph.text not in ('\n', ''):
                        if paragraph.text == '单元测评小结':
                            self.is_after_chapter4 = True
                            self.is_chapter4 = False
                        else:
                            # 跳过本章第一段落，由‘物理安全’开始循环每一小节
                            if paragraph.text == '物理安全':
                                self.cycle_flag = True
                            if self.cycle_flag:
                                # 标记循环序号
                                r = self.paragraph_ro1.search(paragraph.text)
                                if r:
                                    self.step_count += 1
                                    self.unsatisfied_table_id = r.group()[:-1] + '2'
                                elif self.step_count == 0 or self.paragraph_ro2.match(paragraph.text):
                                    self.step_count += 1
                                else:  # 遇到非预期的段落内容，则认为新小节开始，循环序号置零
                                    self.step_count = 0
                                if self.step_count == 0:  # 取出新小节的段落内容作为小节标题
                                    self.section_name = paragraph.text
                            # print(str(self.step_count) + paragraph.text)
                            if paragraph.text == '结果汇总':  # 设定载入标记，下一个表格即将被表格过程调用load_table载入
                                self.load_table_flag = True

                            elif paragraph.text == '结果分析':  # 设定载入标记，下一个表格即将被表格过程调用load_table载入
                                ############################
                                # 生成结果分析文字部分
                                have_unsatisfied = False
                                for b in self.table_some_and_not_satisfy_values:
                                    have_unsatisfied = have_unsatisfied or b
                                have_satisfied = False
                                for b in self.table_satisfy_values:
                                    have_satisfied = have_satisfied or b
                                # print(have_satisfied, have_unsatisfied)
                                if have_unsatisfied:
                                    # 如果有不符合项时，则文档内有不符合项表4-x-2
                                    satisfied_list = list(filter(lambda x: x != '',
                                                                 map(lambda x, y, z: x if y and not z else '',
                                                                     self.table_col_names,
                                                                     self.table_satisfy_values,
                                                                     self.table_some_and_not_satisfy_values)))
                                    unsatisfied_list = list(filter(lambda x: x != '',
                                                                   map(lambda x, y: x if y else '',
                                                                       self.table_col_names,
                                                                       self.table_some_and_not_satisfy_values)))
                                    # print(self.unsatisfied_table_id)
                                    # print(self.section_name, satisfied_list, unsatisfied_list)
                                    if len(satisfied_list) > 0:
                                        output_str1 = '符合项分析：物理安全除表{}所列项目之外，在{}方面均符合要求，在{}方面具备一定的安全性。'.format(
                                            self.unsatisfied_table_id, '、'.join(satisfied_list) + (
                                                '{}'.format('等' if len(satisfied_list) > 1 else '')), self.section_name)
                                    else:
                                        output_str1 = '应用安全在各方面均存在一些问题。'
                                    output_str2 = '部分符合和不符合项分析：{}在{}方面还存在一些问题，详见表{}。'.format(self.section_name,
                                                                                             '、'.join(
                                                                                                 unsatisfied_list) + (
                                                                                                 '{}'.format('等' if len(
                                                                                                     unsatisfied_list) > 1 else '')),
                                                                                             self.unsatisfied_table_id)
                                    print(output_str1)
                                    print(output_str2)
                                    opp = self.insert_paragraph_after(paragraph, output_str1, 'NER-CONTENTTEXT3')
                                    self.insert_paragraph_after(opp, output_str2, 'NER-CONTENTTEXT3')
                                elif not have_unsatisfied and have_satisfied:
                                    # 如果没有不符合项时，则文档内无不符合项表
                                    satisfied_list = self.table_col_names
                                    output_str = '在{}方面均符合要求，在{}方面具备一定的安全性。'.format('、'.join(self.table_col_names) + (
                                        '{}'.format('等' if len(satisfied_list) > 1 else '')),
                                                                                    self.section_name)
                                    print(output_str)
                                    self.insert_paragraph_after(paragraph, output_str, 'NER-CONTENTTEXT3')
                                else:
                                    pass
                                    print('-' * 20)
                                ############################

        def insert_paragraph_after(self, paragraph, text=None, style=None):
            """Insert a new paragraph after the given paragraph."""
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            if style is not None:
                new_para.style = style
            return new_para

        def load_table(self, table):
            self.table_satisfy_values = [False for x in range(15)]
            self.table_some_and_not_satisfy_values = [False for x in range(15)]
            if self.load_table_flag:
                # 保存“结果汇总表”的数据
                for i, row in enumerate(table.rows):
                    try:
                        if i == 1:
                            self.table_col_names = [cell.text for cell in row.cells[3:]]
                        elif (i - 2) % 4 == 0:
                            if len(self.table_satisfy_values) > 0:
                                self.table_satisfy_values = list(map(lambda x, y: x or y, self.table_satisfy_values,
                                                                     [bool(int(cell.text)) if i > 2 else
                                                                      cell.text for i, cell in
                                                                      enumerate(row.cells)][3:]))

                        elif i > 1 and ((i - 2) % 4 == 1 or (i - 2) % 4 == 2):
                            if len(self.table_some_and_not_satisfy_values) > 0:
                                self.table_some_and_not_satisfy_values = list(map(lambda x, y: x or y,
                                                                                  self.table_some_and_not_satisfy_values
                                                                                  , [bool(int(cell.text)) if i > 2 else
                                                                                     cell.text for i, cell in
                                                                                     enumerate(row.cells)][3:]))
                        else:
                            pass
                    except Exception as e:
                        print(i, e)
                # print(self.table_col_names)
                # print(self.table_satisfy_values)
                # print(self.table_some_and_not_satisfy_values)
            self.load_table_flag = False  # 重置标记，等待段落过程重新调用__call__中的设置标志

    # ----------------------------- Cant split Task
    class TableCantsplitTask(Task):
        def __init__(self):
            super(DocxProcessor.TableCantsplitTask, self).__init__()

        def __call__(self, table):
            for row in table.rows:
                tp = parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w')))
                if row._tr.trPr is None:  # trPr为None的情况直接添加属性
                    # print(tp, '--added')
                    row._tr.get_or_add_trPr().append(tp)
                else:  # trPr不为None的情况下才做属性遍历，判断该属性是否已经存在
                    for p in row._tr.trPr.getchildren():
                        # print(p, '--discovered')
                        if 'cantSplit' not in p.tag:
                            row._tr.get_or_add_trPr().append(tp)

    # ----------------------------- Chapter6 Task
    class TableChapter6Task(Task):
        def __init__(self):
            super(DocxProcessor.TableChapter6Task, self).__init__()

        def __call__(self, table):
            pass

    # ============================= DocxProcessor
    def __init__(self, processor_cfg, docx_file):
        assert isinstance(processor_cfg, Config)
        self.file = docx_file
        self.cfg = processor_cfg
        self.document = None
        try:
            self.document = Document(docx_file)
        except PackageNotFoundError as e:
            print(e)
            exit(0)

        self.paragraph_replace_task = DocxProcessor.ParagraphReplaceTask(processor_cfg)
        self.paragraph_chapter4_task = DocxProcessor.ParagraphChapter4Task(processor_cfg)
        self.table_cantsplit_task = DocxProcessor.TableCantsplitTask()
        self.table_chapter6_task = DocxProcessor.TableChapter6Task()

        self.paragraph_process_list = [self.paragraph_replace_task if self.cfg.replace_enabled else None,
                                       self.paragraph_chapter4_task if self.cfg.chapter4_enabled else None
                                       ]
        self.table_process_list = [self.table_cantsplit_task if self.cfg.table_cant_split else None,
                                   self.table_replace_func if self.cfg.replace_enabled else None,
                                   self.paragraph_chapter4_task.load_table if self.cfg.chapter4_enabled else None,
                                   self.table_chapter6_task if self.cfg.chapter6_enabled else None
                                   ]

    def start(self):  # 执行文档处理过程，处理结束后保存
        for block in self.iter_block_items(self.document):
            if isinstance(block, Paragraph):
                self.process_paragraph(block)

            elif isinstance(block, Table):
                self.process_table(block)

        inst_options = [{'selector': '1', 'prompt': '确认保存', 'return': False},
                        {'selector': '2', 'prompt': '退出/不保存', 'return': True}]
        puts(colored.yellow('=' * 50))
        quit = prompt.options('以上内容将发生改变，请确认：', inst_options)
        if quit:
            puts(colored.green('变更未作保存。'))
            exit(0)
        while True:
            try:
                self.document.save(self.file)  # 保存处理完毕的文件
                puts(colored.red('变更已保存。'))
                exit(0)
            except PermissionError as e:
                print(e)
                inst_options = [{'selector': '1', 'prompt': '重试', 'return': False},
                                {'selector': '2', 'prompt': '退出/不保存', 'return': True}]
                quit = prompt.options('文件是否已经打开？请关闭后重试。', inst_options)
                if quit:
                    exit(0)

    def process_paragraph(self, paragraph):

        for func in self.paragraph_process_list:  # 顺次以paragraph为参数，执行列表中的函数
            if func is not None:
                func(paragraph)
        #
        # if 'a' in paragraph.text:
        #     print('in')
        #     for i in paragraph.runs:
        #         if 'a' in i.text:
        #             i.text = i.text.replace('a', 'd')

    def process_table(self, table):
        for func in self.table_process_list:
            if func is not None:
                func(table)

    def table_replace_func(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.paragraph_replace_task(paragraph)

    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """

    def iter_block_items(self, parent):
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        elif isinstance(parent, _Row):
            parent_elm = parent._tr
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
